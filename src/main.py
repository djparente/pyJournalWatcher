# Daniel J. Parente, MD PhD
# University of Kansas Medical Center

from oai import summary_from_cache_or_create
from pymed import PubMed
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import datetime
import logging
import sys
import markdown
from gooey import Gooey, GooeyParser
import toml
from configuration import config
from globalconf import globalconf
import os

# Set up logging
# Make sure the logging directory exists
if not os.path.exists(globalconf.LOGDIRECTORY):
    os.makedirs(globalconf.LOGDIRECTORY)
logging.basicConfig(filename=os.path.join(globalconf.LOGDIRECTORY,'pubmed-api.log'),
                    level=logging.DEBUG,
                    format='%(asctime)s:%(levelname)s:%(message)s')
root = logging.getLogger()
root.setLevel(logging.DEBUG)
handler = logging.StreamHandler(sys.stdout)
handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
root.addHandler(handler)

def executeMain(conf):
    print(f'Got configuration: {conf}')

    pmid_file = os.path.join(conf.BASEDIR, globalconf.PMID_FILE)
    backup_prefix = globalconf.BACKUP_PREFIX
    outsuffix = globalconf.OUTSUFFIX

    # Ensure the appropriate directories exist for output and backup
    needed_directories = [globalconf.OUTPUT_DIRECTORY, globalconf.BACKUP_DIRECTORY]
    needed_directories = [conf.BASEDIR] + [os.path.join(conf.BASEDIR, x) for x in needed_directories]
    for needed_dir in needed_directories:
        if not os.path.exists(needed_dir):
            os.makedirs(needed_dir)

    # Get a pubmed object
    pubmed = PubMed(tool=globalconf.NLM_TOOL_NAME, email=globalconf.NLM_EMAIL)

    # Get the query
    query = conf.QUERY

    # Get the curernt time and date
    nowstr = datetime.datetime.now().isoformat().replace(":", '-').replace('.','_')

    # Execute the query against the API
    results = pubmed.query(query, max_results=conf.MAX_RESULTS, reldate=conf.RELDATE)

    # Keep track of how many are new as we iterate through the file
    new = 0

    # Make sure the pmid file exists
    with open(pmid_file, 'a') as f:
        pass

    # Make a copy of the processed pmids file (as a backup)
    with open(pmid_file, 'r') as f:
        backup_filename = os.path.join(conf.BASEDIR, globalconf.BACKUP_DIRECTORY, f'{backup_prefix}{nowstr}.txt.bak')
        with open(backup_filename, 'w') as bak:
            for x in f.readlines():
                print(x.replace("\r\n", "").replace("\n", ""), file=bak)

    # Get a list of preciously processed PMIDs
    seen_pmids = {}
    with open(pmid_file, 'r') as f:
        seen_pmids = { x.replace("\r\n", "").replace("\n","") : 1 for x in f.readlines() }

    # Create new (docx) document
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)

    # Already create lists of lines for the markdown and markdown simple output formats
    markdown_lines = []
    markdown_lines_simple = []

    # Transform results (an enumerable) into a list, so we can use list comprehensions to process it
    results = [ x for x in results]

    # Filter out articles that we've already seen or that have null abstracts (and thus will be skipped for now)
    already_seen_list = [x for x in results if x.pubmed_id in seen_pmids]
    skippable_list = [x for x in results if (x.abstract is None and x.pubmed_id not in seen_pmids) ]

    # Everything remaining we will need to potentially process
    remaining = [x for x in results if x not in already_seen_list and x not in skippable_list]

    # Keep track of this information in a log
    logging.info(f'Total results: {len(results)}')
    logging.info(f'Already seen: {len(already_seen_list)}')
    logging.info(f'Skippable: {len(skippable_list)}')
    logging.info(f'New: {len(remaining)}')

    already_seen = len(already_seen_list)
    skipped = len(skippable_list)

    # Estimate the cost
    cost_per_1k = 0.002 if conf.GPT_MODEL == "gpt-3.5-turbo" else 0.06
    estimate_cost = len(remaining) * 800 * cost_per_1k / 1000.0
    # Hard stop: If the cost exceeds a maximum, bail out
    if estimate_cost > globalconf.MAX_COST:
        logging.info('Hardcoded maximum estimated cost limit exceeded. Bailing out. Consider reducing lookback period or narrowing query.')
        exit()

    # This used to be an interactive discussion with the user; commented out for now
    #ok_to_proceed = input(f'Planning to summarize {len(remaining)} abstracts using {config.GPT_MODEL} at estimated cost ${estimate_cost}. Proceed?')
    #if ok_to_proceed != "Y":
    #    logging.info('Bailing out per user request')
    #    exit()

    # Open the PMID file to append processed pmids to
    with open(pmid_file, 'a') as f:
        # For each article in the file...
        for i, article in enumerate(remaining):
            # Extract and format information from the article
            article_id = article.pubmed_id
            title = article.title
            journal = article.journal
            doi = article.doi
            publication_date = article.publication_date
            abstract = article.abstract
            sabstract = article.structuredAbstract

            # Calculate the author string
            authstrings = []
            for author in article.authors:
                lname = author['lastname']
                fname_docx = author['firstname']
                nameinit = author['initials']
                authstr = f'{lname}, {fname_docx} {nameinit}'
                authstrings.append(authstr)
            authors = "; ".join(authstrings)

            # We're processing this article, so it's a "new" article in our output
            new += 1

            # If this isn't the first article in our file, the DOCX needs a pagebreak before each article
            if new > 1:
                document.add_page_break()

            # Compile the abstract and markdown abstract
            abstract_plain_paragraphs = []  # The abstract in plain text
            abstract_md_paragraphs = []     # The abstract in markdown text
            if sabstract is not None:       # If there is a structured abstract
                for abspara in sabstract:   # For each item in the structured abstract
                    plainpara = ""          # Accumuator for the current line (plain text)
                    mdpara = ""             # Accumulator for the current line (markdown text)
                    if abspara[0] != "":    # If the structured abstract has a heading
                        plainpara += f'{abspara[0]}: '  # Annotate the heading (plain text), e.g., "Methods"
                        mdpara += f'**{abspara[0]}**: ' # Annotate the heading (markdown)
                    plainpara += abspara[1]             # Then add the body for this part of the abstract (e.g.,
                                                        # "We conducted a multicenter randomized controlled..."
                                                        # (plain text format)
                    mdpara += abspara[1]                # Body for this part of the abstract in markdown
                    abstract_plain_paragraphs.append(plainpara) # Append this line to the growing abstract (plain text)
                    abstract_md_paragraphs.append(mdpara)       # Append this line to the growing abstract (markdown)

            abstract_plain = "\n\n".join(abstract_plain_paragraphs) # Concatenate the abstract lines with newlines between sections
            abstract_md = "\n\n".join(abstract_md_paragraphs)       # And concatenate for the markdown format

            # Prepare to obtain an OpenAI summary
            oai_summary = None
            # Only obtain an OpenAI summary if the abstract has enough characters
            # (no sense in 'summarizing' a 100-character abstract
            if len(abstract_plain) > globalconf.OAI_LOWER_THRESHOLD:
                # Send a call to summarize this article (or pull it from a cache of prior abstracts
                # (the ache pull is based on article_id, not the abstract text)
                oai_summary = summary_from_cache_or_create(
                    article_id,
                    abstract_plain,
                    conf.API_KEY,
                    model=conf.GPT_MODEL,
                    baseDirectory=conf.BASEDIR,
                    simple_instructions=False
                )

            # Output in DOCX format
            p = document.add_paragraph()
            p.add_run(article.title).bold = True
            document.add_paragraph(authors)

            # In the DOCX file, create a section for users to enter their own free-text notes on the article
            # while reviewing it
            usernotes = document.add_paragraph()                    # Create a paragraph
            usernotes.paragraph_format.left_indent = Inches(0.5)    # Indent that paragraph
            usernotesimprun = usernotes.add_run("Importance:")      # Call this "Importance" for the user
            usernotesimprun.bold = True                             # Heading is bold
            usernotesimprun.font.color.rgb = RGBColor(255, 0, 0)    # Heading is red
            usernotesimprun = usernotes.add_run(" ***")             # Create *** as a placeholder for user remarks
            usernotesimprun.font.color.rgb = RGBColor(255, 0, 0)    # User remarks will also be red

            # If we have a summary of the article, we will also want to append this to the DOCX
            if oai_summary is not None:
                gptnotes = document.add_paragraph()                 # Create new paragraph
                gptnotes.paragraph_format.left_indent = Inches(0.5) # Indent the summary
                gptimprun = gptnotes.add_run(f'{conf.GPT_NAME} Summary: ')  # Heading for the summary
                gptimprun.bold = True                                       # Heading is bold
                gptimprun.font.color.rgb = RGBColor(0, 0, 255)              # Heading is blue
                gptimprun2 = gptnotes.add_run(oai_summary)                  # Add the summary text
                gptimprun2.font.color.rgb = RGBColor(0, 0, 255)             # Summary text is also blue

            # If there is an abstract, we'll need to add this as well
            if sabstract is not None:
                # For each section in the abstract (Introduction, Methods, etc.)...
                # (For unstructured abstracts there is a single section with no heading)
                for abspara in sabstract:
                    ap = document.add_paragraph()           # Add a paragraph to the documenbt
                    if abspara[0] != "":                    # If this part of the abstract has a heading (e.g. Methods)
                        ap.add_run(abspara[0]).bold = True  #   add that heading and make it bold
                        ap.add_run(": ")                    #   and also put a colon after it

                    ap.add_run(abspara[1])                  # Add the abstract section text (e.g., "We conducted a ...")

            # Add metadata to the footer of the article
            document.add_paragraph(f'{publication_date} - {journal} - {article_id} - {doi}')

            # Writing this in markdown and plaintext format is much fewer lines, do that in the two blocks below

            # Write markdown format
            markdown_lines.append(f'\n')                    # New line for new article
            markdown_lines.append(f'## {article.title}')    # Article title
            markdown_lines.append(f'{authors}')             # Author string
            if oai_summary is not None:                     # Add the GPT summary, if it exists
                markdown_lines.append(f'\n### {conf.GPT_NAME} Summary: ')   # Header line
                markdown_lines.append(f'{oai_summary}')                     # Body of the summary
            markdown_lines.append('\n### Abstract')         # Add the abstract header
            markdown_lines.append(abstract_md)              # Add the abstract itself
            # Add metadata (including hyperlinks!)
            markdown_lines.append(f'\n{publication_date} - {journal} - [{article_id}](https://pubmed.ncbi.nlm.nih.gov/{article_id}) - [{doi}](https://dx.doi.org/{doi})')

            # And in the abbreviated simple format (which omits the abstracts, unless there is no GPT summary, typically
            # because the article is too short, in which case it includes the full text of the abstract)
            markdown_lines_simple.append(f'\n')                 # New line for new article
            markdown_lines_simple.append(f'## {article.title}') # Article title
            markdown_lines_simple.append(f'{authors}')          # Author string
            # If there is an GPT summary, show (only) it, otherwise show the abstract full text
            if oai_summary is not None:
                # GPT summary
                markdown_lines_simple.append(f'\n### {conf.GPT_NAME} Summary') # GPT summary header line
                markdown_lines_simple.append(f'{oai_summary}')                 # Summary text
            else:
                markdown_lines_simple.append('\n### Abstract')                 # Abstract summary header line
                markdown_lines_simple.append(abstract_md)                      # Abstract summary body text
            # And add metadata (including hyperlinks)
            markdown_lines_simple.append(f'\n{publication_date} - {journal} - [{article_id}](https://pubmed.ncbi.nlm.nih.gov/{article_id}) - [{doi}](https://dx.doi.org/{doi})')

            # Save this to the PMID file, so that we know we reviewed and output this file
            print(article_id, file=f)

    # Print some output statistics for the user
    print(f'New {new}')
    print(f'Skipped {skipped}')
    print(f'Already seen {already_seen}')

    # If there were any articles summarized at all...
    if new > 0:
        # Get a base file name to output, based on the time we ran the program
        fname_prefix = os.path.join(conf.BASEDIR, globalconf.OUTPUT_DIRECTORY, f'AbstractReview_{nowstr}{outsuffix}')

        # Save the DOCX file
        fname_docx = f'{fname_prefix}.docx'     # Calcualte the filename
        print(f'Writing file {fname_docx}')     # Alert the user
        document.save(fname_docx)               # Save the DOCX file

        # Write the markdown formatted file (note this MUST occur before the HTML file write below, because that uses
        # this file to create the HTML file)
        markdown_filename = fname_prefix + '.md'    # Calculate the filename
        print(f'Writing file {markdown_filename}')  # Alert the user
        with open(markdown_filename, 'w', encoding='utf-8') as mdfile:  # Write the file out (encoding matters to avoid errors)
            mdfile.write("\n".join(markdown_lines[1:])) # Write every line of the markdown file except the first one (which is an extraneous newline before the first abstract)

        # Likewise, write the 'simple' file that includes only the GPT summaries, and not the abstracts (unless the
        # abstracts were too short to require a summary, in which case the abstracts are included)
        mdnamesimple = fname_prefix + '_simple.md'  # Calculate filename
        print(f'Writing file {mdnamesimple}')  # Alert the user
        with open(mdnamesimple, 'w', encoding='utf-8') as mdfile:   # Write the file out (again, encoding matters to avoid errors)
            mdfile.write("\n".join(markdown_lines_simple[1:]))      # Again write out every line except the first line (an extraneous newline)

        # Convert the markdown file to HTML
        htmlname = fname_prefix + '.html'
        with open(markdown_filename, 'r', encoding='utf-8') as mdfile:
            # User the markdown package to convert the markdown format into an HTML format
            html = markdown.markdown(mdfile.read()).replace("\r\n", "\n")

            # The HTML file needs to have a header that specifies the UTF-8 encoding (otherwise encoding errors are
            # clearly evident throughout), so add this header. Because there is a header, wrap the HTML in body tags
            html = f'<head><meta charset="UTF-8"></head>\n<body>\n{html}\n</body>'

            print(f'Writing file {htmlname}')  # Alert the user
            with open(htmlname, 'w', encoding="utf-8") as htmlfile: # Write the HTML file, again encoding is important
                print(html, file=htmlfile)  # Why do we use print instead of .write()? We just do...
    else:
        print("No updates to write") # If there weren't any updates, instead tell the user we didn't write any files

    # That's it. Program complete.


@Gooey(
    program_name='pyJournalWatch',       # Defaults to script name
    program_description='Rapid and systematic surveillance of the biomedical literature augmented by artificial intelligence',       # Defaults to ArgParse Description
    default_size=(900, 1000),   # starting size of the GUI
    required_cols=1,           # number of columns in the "Required" section
    optional_cols=2,           # number of columns in the "Optional" section
    dump_build_config=False,   # Dump the JSON Gooey uses to configure itself
    load_build_config=None,    # Loads a JSON Gooey-generated configuration
    monospace_display=False   # Uses a mono-spaced font in the output screen
)
def main():
    lastguiconf_path = os.path.join(globalconf.DATADIR, 'lastguiconf.toml')
    lastgui = {}
    if os.path.exists(lastguiconf_path):
        try:
            with open(lastguiconf_path, 'r') as lastconf:
                lastgui = toml.load(lastconf)
        except:
            logging.warning('No prior GUI configuration file found')

    missing_api_key = 'xxx-obtain-api-key-from-openai'
    use_last_saved_api_key = '<Use last saved API key>'

    default_api_key = missing_api_key
    saved_api_key = missing_api_key
    if 'API_KEY' in lastgui:
        saved_api_key=lastgui['API_KEY']
        default_api_key = use_last_saved_api_key

    default_lookback = 7
    if 'RELDATE' in lastgui:
        default_lookback = lastgui['RELDATE']

    default_writtenquery = ''
    if 'WRITTENQUERY' in lastgui:
        default_writtenquery = lastgui['WRITTENQUERY']

    default_basedir = ''
    if 'BASEDIR' in lastgui:
        default_basedir = lastgui['BASEDIR']


    parser = GooeyParser(description="A program for systemic surveillance of the medical literature")
    #config_group = parser.add_argument_group("Configuration file",
    #                                         "Optionally select a configuration file instead of specifying these settings")
    #config_group.add_argument("--config", widget='FileChooser',
    #                    metavar="Configuration File",
    #                    help="If specified, will override all these settings",
    #                    default='')

    basic_options = parser.add_argument_group("Basic Options",
                                              "Specify the lookback period and the API key")

    basic_options.add_argument("--lookback", type=str, default=default_lookback, metavar="Lookback period",
                        help="Number of days to look back")
    basic_options.add_argument("--api_key", type=str,
                        metavar="OpenAI.com API Key",
                        help="Can be obtained from OpenAI.com",
                        default=default_api_key)
    basic_options.add_argument('--output_dir',
                               widget="DirChooser",
                               metavar="Output Directory",
                               gooey_options={'initial_value': default_basedir},
                               help="Directory to store output files, backup and GPT cache")
    basic_options.add_argument('--model',
                               choices=['gpt-3.5-turbo', 'gpt-4'],
                               default='gpt-3.5-turbo',
                               help="Select GPT-3.5 or GPT-4 model to use",
                               metavar="GPT model")

    journal_group = parser.add_argument_group(
        "Common Journals",
        "Select from common family medicine journals"
    )




    journals = [
        ('afm', 'Annals of Family Medicine', 'Annals of family medicine'),
        ('jabfm', 'Journal of the American Board of Family Medicine', 'Journal of the American Board of Family Medicine : JABFM'),
        ('fm_stfm', 'Family Medicine', 'Family medicine"[Journal]'),
        ('jama', 'Journal of the American Medical Association (JAMA)', 'JAMA'),
        ('jama_im', 'JAMA Internal Medicine', 'JAMA internal medicine'),
        ('aim', 'Annals of Internal Medicine', 'Annals of internal medicine'),
        ('nejm', 'New England Journal of Medicine', 'The New England journal of medicine'),
        ('nm', 'Nature Medicine', 'Nature medicine'),
        ('jgim', 'Journal of General Internal Medicine', 'Journal of general internal medicine')
    ]

    guijournals = []
    if 'JOURNALS' in lastgui:
        guijournals = lastgui['JOURNALS']

    for abbrev, jname, pmname in journals:
        abbrev_nice = abbrev.upper().replace("_", " ")

        default_checking = (pmname in guijournals)

        journal_group.add_argument(f'--{abbrev}',
                                   help=f'Include {abbrev_nice}',
                                   metavar=f'{jname}',
                                   action='store_true',
                                   default=False,
                                   gooey_options={'initial_value': default_checking}
                                   )


    query_group = parser.add_argument_group(
        "Ad hoc Pubmed Queries",
    )

    query_group.add_argument('--query',
                             type=str,
                             metavar='Query string',
                             help="Additional queries can be specified in PubMed's search format",
                             gooey_options={'initial_value': default_writtenquery},
                             default=''
                             )

    args = parser.parse_args()

    journallist = [x[2] for x in journals if getattr(args, x[0])]
    journalquery = config.build_journal_query(journallist)

    fullQuery = journalquery
    if args.query is not None and args.query != '':
        if fullQuery == '':
            fullQuery = args.query
        else:
            fullQuery = config.concat_queries(args.query, journalquery)

    if fullQuery == '':
        print("Error: No journals or queries specified; bailing out")
        exit(-1)

    api_key = args.api_key
    if api_key is None or api_key == 'xxx-obtain-api-key-from-openai':
        print("Error: No API key specified. Obtain one from OpenAI.com and specify it here.")
        exit(-1)
    elif api_key == use_last_saved_api_key:
        api_key = saved_api_key

    reldate = args.lookback
    if reldate is None:
        print("Error: No lookback period specified.")
        exit(-1)

    output_dir = args.output_dir
    if output_dir is None:
        print("Error: No output directory")
        exit(-1)

    gpt_model = "gpt-3.5-turbo"
    gpt_model_name = "GPT-3.5"
    if args.model is not None and args.model == 'gpt-4':
        gpt_model = "gpt-4"
        gpt_model_name = "GPT-4"

    conf = config(
        apikey=api_key,
        gpt_model=gpt_model,
        gpt_model_name=gpt_model_name,
        max_results=1000,
        reldate=reldate,
        query=fullQuery,
        writtenquery=args.query,
        journals=journallist,
        basedir=output_dir
    )

    # Save last known GUI configuration

    conf.to_toml(lastguiconf_path)
    executeMain(conf)


if __name__ == "__main__":
    main()
